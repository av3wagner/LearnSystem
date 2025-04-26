#Stand: 20.11.2023 
import streamlit as st  
import os, sys
from os import listdir
from os.path import isfile, join
import pathlib
import base64
from docx import Document

def show_pdf(file_path):
    st.title('✨ Визуализация PDF документа 📜')
    st.markdown("")
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="1000" height="700" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

#import time
#exec(open("/opt/render/project/src/modules/programs/MLReports.py").read(), globals())
#time.sleep(2)
#st.write('Ende Programm!')

def execute_python_file1(file_path):
    try:
        import time
        #exec(open(file_path).read(), globals())
        exec(open("modules/programs/EDAReports.py").read(), globals())
        time.sleep(2)
    except FileNotFoundError:
        st.markdown(f"Error: The file '{file_path}' does not exist.")

def execute_python_file(file_path):
    #file_path="modules/programs/EDAReports.py"
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            python_code = file.read()
            import time
            time.sleep(3)
            exec(python_code)
            time.sleep(2)
    except FileNotFoundError:
        st.markdown(f"Error: The file '{file_path}' does not exist.")

def select_file():
    #parent_path = 'modules\programs'
    parent_path = '.\modules\programs'
    fileList = []
    extensions = ['py']
    fileList = listdir(parent_path)
    onlyfiles = [f for f in fileList if isfile(join(parent_path, f)) and  (f.endswith(".py"))]   
    option = st.selectbox('Выберите программу для EDA/ML-Анализа', onlyfiles)
    file_location=os.path.join(parent_path, option) 
    if file_location.find('.py') > 0:
        if st.button('Запустите EDA/ML-программу'):
            st.write(file_location) 
            execute_python_file(file_location)
            
        if st.button('Покажите EDA/Ml-программу'):    
            with open(file_location, 'r', encoding='utf-8') as f:
                 lines_to_display = f.read()
            st.code(lines_to_display, "python")    

def open_file_selection_doc():
    parent_path = 'modules/docs'
    fileList = []
    extensions = ['pdf', 'docx']
    fileList = listdir(parent_path)
    onlyfiles = [f for f in fileList if isfile(join(parent_path, f)) and  (f.endswith(".pdf") or f.endswith(".docx"))]   
    option = st.selectbox('Выберите Документ', onlyfiles)
    file_location=os.path.join(parent_path, option) 
    
    if file_location.find('.pdf') > 0:
         if st.button('Покажите Документ'):    
            #st.write(file_location) 
            #with st.expander("1. PDF Документ", expanded=True):
            show_pdf(file_location)
      
    elif file_location.find('.docx') > 0:
         if st.button('Покажите Документ'):    
            st.write(file_location) 
            doc = Document(file_location)
            all_paras = doc.paragraphs
            for para in all_paras:
                #print(para.text)   
                #st.code(para.text, "python")
                st.write(para.text) 
    
            
def open_file_selection_doc2():
    parent_path = 'modules/docs'
    fileList = []
    extensions = ['pdf', 'docx']
    fileList = listdir(parent_path)
    onlyfiles = [f for f in fileList if isfile(join(parent_path, f)) and  (f.endswith(".pdf") or f.endswith(".docx"))]   
    option = st.selectbox('Выберите Документ', onlyfiles)
    file_location=os.path.join(parent_path, option) 
    
    if file_location.find('.pdf') > 0:
         if st.button('Покажите Документ'):    
            st.write(file_location) 
            reader = PdfFileReader(file_location)
            no_pages = reader.numPages
            i = 0
            while i < no_pages:
                page = reader.pages[i]
                #print(page.extract_text())
                st.code(page, "python") 
                i += 1 
    elif ile_location.find('.docx') > 0:
            doc = Document(file_location)
            all_paras = doc.paragraphs
            for para in all_paras:
                #print(para.text)   
                st.code(para.text, "python")
                
def open_test():    
    filenames = fd.askopenfilenames()
    for filename in filenames:
        extension = pathlib.Path(filename).suffix
        if extension == '.pdf':
            st.write(filename) 
            reader = PdfFileReader(filename)
            no_pages = reader.numPages
            i = 0
            while i < no_pages:
                page = reader.pages[i]
                print(page.extract_text())
                i += 1
        elif extension == '.txt':
            with open(filename, 'r') as f:
                read_data = f.read()
                print(read_data)
        elif extension in ['.doc', '.docx']:
            doc = Document(filename)
            all_paras = doc.paragraphs
            for para in all_paras:
                print(para.text)
        else:
            print("Can't read files with extension {} for file {}".format(extension, filename))



  